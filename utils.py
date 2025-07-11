import io
import re
import logging
from typing import List, Optional
import PyPDF2
import pdfplumber
from docx import Document
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from various file formats
    
    Args:
        file_content: Binary content of the file
        filename: Name of the file to determine format
        
    Returns:
        Extracted text as string
    """
    
    try:
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return extract_text_from_pdf(file_content)
        elif file_extension in ['docx', 'doc']:
            return extract_text_from_docx(file_content)
        elif file_extension == 'txt':
            return extract_text_from_txt(file_content)
        else:
            logging.warning(f"Unsupported file format: {file_extension}")
            return ""
            
    except Exception as e:
        logging.error(f"Error extracting text from {filename}: {str(e)}")
        return ""

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    
    text = ""
    
    try:
        # Try with pdfplumber first (better for complex layouts)
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # If pdfplumber didn't work well, try PyPDF2
        if len(text.strip()) < 100:
            text = ""
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
    except Exception as e:
        logging.error(f"Error extracting PDF text: {str(e)}")
        
        # Fallback to PyPDF2 if pdfplumber fails
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        except Exception as e2:
            logging.error(f"Fallback PDF extraction also failed: {str(e2)}")
    
    return clean_extracted_text(text)

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return clean_extracted_text(text)
        
    except Exception as e:
        logging.error(f"Error extracting DOCX text: {str(e)}")
        return ""

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file"""
    
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                return clean_extracted_text(text)
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, use utf-8 with error handling
        text = file_content.decode('utf-8', errors='ignore')
        return clean_extracted_text(text)
        
    except Exception as e:
        logging.error(f"Error extracting TXT text: {str(e)}")
        return ""

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep important punctuation
    text = re.sub(r'[^\w\s@.,()-]', ' ', text)
    
    # Remove multiple spaces
    text = re.sub(r' +', ' ', text)
    
    # Remove leading/trailing whitespace
    text = text.strip()
    
    return text

def extract_skills_from_text(text: str) -> List[str]:
    """
    Extract technical skills from text using keyword matching
    
    Args:
        text: Input text to analyze
        
    Returns:
        List of identified skills
    """
    
    # Comprehensive list of technical skills
    technical_skills = [
        # Programming Languages
        'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust',
        'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'Perl', 'Shell', 'Bash', 'PowerShell',
        
        # Web Technologies
        'HTML', 'CSS', 'React', 'Angular', 'Vue.js', 'Node.js', 'Express.js', 'Django', 'Flask',
        'Spring', 'Laravel', 'Ruby on Rails', 'ASP.NET', 'jQuery', 'Bootstrap', 'Sass', 'Less',
        'Webpack', 'Gulp', 'Grunt', 'Next.js', 'Nuxt.js', 'Gatsby',
        
        # Databases
        'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'SQLite', 'Oracle', 'SQL Server', 'Cassandra',
        'DynamoDB', 'Neo4j', 'Elasticsearch', 'InfluxDB', 'CouchDB', 'MariaDB',
        
        # Cloud & DevOps
        'AWS', 'Azure', 'Google Cloud', 'GCP', 'Docker', 'Kubernetes', 'Jenkins', 'GitLab CI',
        'GitHub Actions', 'Terraform', 'Ansible', 'Chef', 'Puppet', 'Vagrant', 'Nginx', 'Apache',
        'CloudFormation', 'Helm', 'Istio',
        
        # Data Science & ML
        'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch', 'Scikit-learn', 'Pandas',
        'NumPy', 'Matplotlib', 'Seaborn', 'Jupyter', 'Apache Spark', 'Hadoop', 'Kafka',
        'Airflow', 'MLflow', 'Kubeflow', 'OpenCV', 'NLTK', 'spaCy',
        
        # Mobile Development
        'iOS', 'Android', 'React Native', 'Flutter', 'Xamarin', 'Ionic', 'Cordova',
        
        # Tools & Frameworks
        'Git', 'SVN', 'JIRA', 'Confluence', 'Slack', 'Trello', 'Asana', 'Figma', 'Sketch',
        'Photoshop', 'Illustrator', 'InDesign', 'Blender', 'Unity', 'Unreal Engine',
        'Visual Studio', 'IntelliJ', 'Eclipse', 'Vim', 'Emacs',
        
        # Testing
        'Selenium', 'Jest', 'Mocha', 'Chai', 'Cypress', 'Puppeteer', 'JUnit', 'TestNG',
        'PyTest', 'RSpec', 'Cucumber',
        
        # Methodologies & Concepts
        'Agile', 'Scrum', 'Kanban', 'DevOps', 'CI/CD', 'TDD', 'BDD', 'Microservices',
        'RESTful APIs', 'GraphQL', 'SOAP', 'OAuth', 'JWT', 'API Gateway', 'Load Balancing',
        'Caching', 'CDN', 'Monitoring', 'Logging'
    ]
    
    found_skills = []
    text_lower = text.lower()
    
    for skill in technical_skills:
        # Use word boundaries for more accurate matching
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, text_lower):
            found_skills.append(skill)
    
    # Remove duplicates and sort
    return sorted(list(set(found_skills)))

def extract_contact_info(text: str) -> dict:
    """
    Extract contact information from text
    
    Args:
        text: Input text to analyze
        
    Returns:
        Dictionary containing contact information
    """
    
    contact_info = {
        'emails': [],
        'phones': [],
        'linkedin': [],
        'github': [],
        'websites': []
    }
    
    # Email extraction
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    contact_info['emails'] = re.findall(email_pattern, text)
    
    # Phone extraction
    phone_patterns = [
        r'\+?1?[-.\s]?$$?([0-9]{3})$$?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
        r'$$\d{3}$$\s?\d{3}[-.\s]?\d{4}'
    ]
    
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple):
                phone = f"({match[0]}) {match[1]}-{match[2]}"
            else:
                phone = match
            contact_info['phones'].append(phone)
    
    # LinkedIn extraction
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'
    contact_info['linkedin'] = re.findall(linkedin_pattern, text, re.IGNORECASE)
    
    # GitHub extraction
    github_pattern = r'github\.com/[\w-]+'
    contact_info['github'] = re.findall(github_pattern, text, re.IGNORECASE)
    
    # Website extraction
    website_pattern = r'https?://(?:[-\w.])+(?:[:\d]+)?(?:/(?:[\w/_.])*(?:\?(?:[\w&=%.])*)?(?:#(?:[\w.])*)?)?'
    contact_info['websites'] = re.findall(website_pattern, text)
    
    # Remove duplicates
    for key in contact_info:
        contact_info[key] = list(set(contact_info[key]))
    
    return contact_info

def calculate_text_statistics(text: str) -> dict:
    """
    Calculate various text statistics
    
    Args:
        text: Input text to analyze
        
    Returns:
        Dictionary containing text statistics
    """
    
    try:
        # Basic statistics
        word_count = len(word_tokenize(text))
        char_count = len(text)
        sentence_count = len(nltk.sent_tokenize(text))
        
        # Average word length
        words = word_tokenize(text.lower())
        avg_word_length = sum(len(word) for word in words) / len(words) if words else 0
        
        # Unique words
        unique_words = len(set(words))
        
        # Stop words ratio
        stop_words = set(stopwords.words('english'))
        stop_word_count = sum(1 for word in words if word in stop_words)
        stop_word_ratio = stop_word_count / len(words) if words else 0
        
        return {
            'word_count': word_count,
            'character_count': char_count,
            'sentence_count': sentence_count,
            'average_word_length': round(avg_word_length, 2),
            'unique_words': unique_words,
            'vocabulary_richness': round(unique_words / word_count, 3) if word_count > 0 else 0,
            'stop_word_ratio': round(stop_word_ratio, 3)
        }
        
    except Exception as e:
        logging.error(f"Error calculating text statistics: {str(e)}")
        return {
            'word_count': 0,
            'character_count': len(text),
            'sentence_count': 0,
            'average_word_length': 0,
            'unique_words': 0,
            'vocabulary_richness': 0,
            'stop_word_ratio': 0
        }

def normalize_text(text: str) -> str:
    """
    Normalize text for better processing
    
    Args:
        text: Input text to normalize
        
    Returns:
        Normalized text
    """
    
    # Convert to lowercase
    text = text.lower()
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters except important ones
    text = re.sub(r'[^\w\s@.,()-]', ' ', text)
    
    # Remove multiple spaces
    text = re.sub(r' +', ' ', text)
    
    # Strip leading/trailing whitespace
    text = text.strip()
    
    return text

def extract_years_from_text(text: str) -> List[int]:
    """
    Extract years from text
    
    Args:
        text: Input text to analyze
        
    Returns:
        List of years found in text
    """
    
    # Pattern to match 4-digit years (1900-2099)
    year_pattern = r'\b(19|20)\d{2}\b'
    years = re.findall(year_pattern, text)
    
    # Convert to integers and remove duplicates
    year_list = sorted(list(set([int(''.join(year)) for year in years])))
    
    return year_list

def validate_email(email: str) -> bool:
    """
    Validate email address format
    
    Args:
        email: Email address to validate
        
    Returns:
        True if valid, False otherwise
    """
    
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return bool(re.match(pattern, email))

def validate_phone(phone: str) -> bool:
    """
    Validate phone number format
    
    Args:
        phone: Phone number to validate
        
    Returns:
        True if valid, False otherwise
    """
    
    # Remove all non-digit characters
    digits_only = re.sub(r'\D', '', phone)
    
    # Check if it's a valid US phone number (10 or 11 digits)
    return len(digits_only) in [10, 11]

def format_phone(phone: str) -> str:
    """
    Format phone number to standard format
    
    Args:
        phone: Phone number to format
        
    Returns:
        Formatted phone number
    """
    
    # Remove all non-digit characters
    digits_only = re.sub(r'\D', '', phone)
    
    if len(digits_only) == 10:
        return f"({digits_only[:3]}) {digits_only[3:6]}-{digits_only[6:]}"
    elif len(digits_only) == 11 and digits_only[0] == '1':
        return f"+1 ({digits_only[1:4]}) {digits_only[4:7]}-{digits_only[7:]}"
    else:
        return phone  # Return original if can't format

def extract_education_keywords(text: str) -> List[str]:
    """
    Extract education-related keywords from text
    
    Args:
        text: Input text to analyze
        
    Returns:
        List of education keywords found
    """
    
    education_keywords = [
        'bachelor', 'master', 'phd', 'doctorate', 'associate', 'diploma', 'certificate',
        'b.s.', 'b.a.', 'm.s.', 'm.a.', 'mba', 'ph.d.', 'b.tech', 'm.tech',
        'university', 'college', 'institute', 'school', 'academy', 'degree',
        'graduation', 'undergraduate', 'graduate', 'postgraduate'
    ]
    
    found_keywords = []
    text_lower = text.lower()
    
    for keyword in education_keywords:
        if keyword in text_lower:
            found_keywords.append(keyword)
    
    return sorted(list(set(found_keywords)))
